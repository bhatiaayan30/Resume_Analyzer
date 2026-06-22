"""
utils.py — business logic for the Resume Analyzer.

Keeping extraction + AI logic here (not in views.py) means
these functions can be tested independently without HTTP.

Usage:
    from analyzer.utils import extract_text, analyze_with_ai
"""

import json
import os
from typing import Dict, Any, Tuple, List

import fitz  # PyMuPDF
import pdfplumber
from docx import Document
from groq import Groq
from tenacity import retry, stop_after_attempt, wait_exponential

from analyzer.prompt_builder import (
    build_analysis_prompt,
    build_cover_letter_system_prompt,
    build_cover_letter_user_prompt,
    build_bullet_rewrite_prompt,
    build_interview_question_prompt,
    build_interview_feedback_prompt,
    build_resume_parser_prompt,
    build_summary_suggestion_prompt,
    build_experience_bullets_prompt,
)
import re

def check_searchability(text: str) -> List[Dict[str, Any]]:
    checks = []
    
    # 1. Contact Info
    has_email = bool(re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text))
    has_phone = bool(re.search(r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text))
    passed_contact = has_email and has_phone
    checks.append({
        "check_name": "Contact Information",
        "passed": passed_contact,
        "detail": "Contact info (email, phone) is present and parseable." if passed_contact else "Missing email or phone number, or it is hidden inside a header/footer region."
    })
    
    # 2. Date pattern
    has_date = bool(re.search(r'(19|20)\d{2}\s*[-–]\s*(19|20)\d{2}|Present', text, re.IGNORECASE))
    checks.append({
        "check_name": "Date Parseability",
        "passed": has_date,
        "detail": "Valid date ranges detected for work experience." if has_date else "No standard date patterns (e.g., 2020 - 2022) found. ATS may fail to parse experience duration."
    })
    
    # 3. Section Headings
    headings = ["experience", "education", "skills", "projects", "summary"]
    found_headings = [h for h in headings if re.search(rf'^\s*{h}\s*$', text, re.MULTILINE | re.IGNORECASE) or re.search(rf'\b{h}\b', text[:500], re.IGNORECASE)]
    passed_headings = len(found_headings) >= 2
    checks.append({
        "check_name": "Standard Section Headings",
        "passed": passed_headings,
        "detail": "Standard section labels (Experience, Education, etc.) used." if passed_headings else "Unconventional section headings detected. ATS might not categorize your experience correctly."
    })
    
    return checks

# ──────────────────────────────────────────────────────────────
# Text extraction
# ──────────────────────────────────────────────────────────────


def extract_text(file_obj: Any, ext: str) -> Tuple[str, List[str], List[Dict[str, Any]]]:
    """
    Extract plain text from an uploaded PDF or DOCX file object.

    Args:
        file_obj : Django InMemoryUploadedFile
        ext      : '.pdf' or '.docx'  (lowercase, includes the dot)

    Returns:
        tuple: (Plain text string, list of ATS format warnings (keys), list of searchability checks)

    Raises:
        ValueError: if the file has no extractable text.
    """
    ats_format_issues = []

    if ext == ".pdf":
        file_obj.seek(0)
        file_bytes = file_obj.read()

        has_tables = False
        has_images = False
        text = ""

        try:
            # Primary robust extraction with PyMuPDF
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            pages = []
            for page in doc:
                # Basic layout-aware extraction
                page_text = page.get_text("text")
                pages.append(page_text)

                # Check for images
                if page.get_images():
                    has_images = True
                    
                # Check for tables
                if hasattr(page, "find_tables") and len(page.find_tables().tables) > 0:
                    has_tables = True

            text = "\n".join(pages).strip()
        except Exception as e:
            raise ValueError(f"Failed to read PDF: {str(e)}")

        if has_tables:
            ats_format_issues.append("tables_detected")
        if has_images:
            ats_format_issues.append("images_detected")

        if not text:
            # Most common user error — scanned / image-based PDF
            raise ValueError(
                "This PDF has no text layer — it's probably a scanned document. "
                "Please upload a text-based PDF, or run it through an OCR tool "
                "(e.g. Adobe Acrobat, Google Drive) and re-export."
            )

        # Garbled text check
        alnum_count = sum(c.isalnum() for c in text)
        if len(text) > 0 and (alnum_count / len(text)) < 0.5:
            raise ValueError(
                "The extracted text appears garbled or heavily formatted with non-standard characters. "
                "The PDF might have font encoding issues."
            )

        return text, ats_format_issues, check_searchability(text)

    elif ext == ".docx":
        doc = Document(file_obj)
        if len(doc.tables) > 0:
            ats_format_issues.append("tables_detected")

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        if not paragraphs:
            raise ValueError("The DOCX file appears to be empty.")
        return "\n".join(paragraphs), ats_format_issues, check_searchability("\n".join(paragraphs))

    raise ValueError(f"Unsupported extension: {ext!r}")

def extract_text_from_image(image_bytes: bytes, mime_type: str) -> str:
    """
    Extract text from an image using Groq Vision API.
    """
    from django.conf import settings
    import base64
    
    api_key = getattr(settings, "GROQ_API_KEY", None) or os.environ.get("GROQ_API_KEY", "")
    if not api_key:
        raise RuntimeError("GROQ_API_KEY environment variable is not set.")

    client = Groq(api_key=api_key)
    base64_encoded = base64.b64encode(image_bytes).decode('utf-8')
    
    try:
        response = client.chat.completions.create(
            model="llama-3.2-11b-vision-preview",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "Extract and transcribe all text from this resume image exactly as it appears. Do not include any conversational filler, markdown formatting blocks, or explanations. Only output the raw text."},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:{mime_type};base64,{base64_encoded}",
                            },
                        },
                    ],
                }
            ],
            temperature=0.1,
            max_tokens=4000,
        )
        return response.choices[0].message.content.strip()
    except Exception as exc:
        raise RuntimeError(f"OCR Failed: {exc}")

# ──────────────────────────────────────────────────────────────
# AI analysis
# ──────────────────────────────────────────────────────────────


@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=2, max=10))
def analyze_with_ai(resume_text: str, job_description: str) -> Tuple[Dict[str, Any], Dict[str, int]]:
    """
    Uses Groq API (Llama-3.1-8b) to score and extract skills/gaps from the text.
    Returns a tuple: (analysis_json, token_usage).

    Returns a dict with keys:
        match_score     : int  (0–100)
        matched_skills  : list[str]
        missing_skills  : list[str]
        experience_gaps : list[str]
        suggestions     : list[str]   — specific, actionable improvements
        upskill_paths   : list[dict]  — suggested learning paths
        impact_critiques: list[dict]  — bullet point writing critiques

    Raises:
        RuntimeError         : GROQ_API_KEY is not set.
        json.JSONDecodeError : AI returned non-JSON.
    """
    from django.conf import settings

    api_key = getattr(settings, "GROQ_API_KEY", None) or os.environ.get("GROQ_API_KEY", "")
    if not api_key:
        raise RuntimeError(
            "GROQ_API_KEY environment variable is not set. " "Add it to your .env file."
        )

    client = Groq(api_key=api_key)


    # ── Fetch Prompt from Builder ─────────────
    prompt = build_analysis_prompt(resume_text, job_description)
    
    fast_model = os.environ.get("FAST_MODEL", "llama-3.3-70b-versatile")

    try:
        response = client.chat.completions.create(
            model=fast_model,
            messages=[
                {"role": "system", "content": "You output strictly valid JSON without markdown wrapping."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.2,
            max_tokens=4000,
            response_format={"type": "json_object"},
        )
        result_text = response.choices[0].message.content
        usage = response.usage
        usage_data = {
            "prompt_tokens": usage.prompt_tokens if usage else 0,
            "completion_tokens": usage.completion_tokens if usage else 0,
        }
        return json.loads(result_text), usage_data
    except Exception as exc:
        raise RuntimeError(f"AI Analysis Failed: {exc}")


@retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=2, max=10))
def generate_cover_letter(resume_text: str, job_desc: str, tone: str = "Professional", length: str = "Medium", highlights: str = "") -> str:
    """
    Uses Groq (Llama-3) to write a highly customized cover letter based on the resume and job description.
    """
    from django.conf import settings

    api_key = getattr(settings, "GROQ_API_KEY", None) or os.environ.get("GROQ_API_KEY", "")
    if not api_key:
        raise ValueError("GROQ_API_KEY is not configured.")

    client = Groq(api_key=api_key)

    system_prompt = build_cover_letter_system_prompt(tone=tone)
    user_prompt = build_cover_letter_user_prompt(resume_text, job_desc, length=length, highlights=highlights)
    
    quality_model = os.environ.get("QUALITY_MODEL", "llama-3.3-70b-versatile")

    response = client.chat.completions.create(
        model=quality_model,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.7,
        max_tokens=1500,
    )

    return response.choices[0].message.content.strip()

def generate_otp() -> str:
    """Generate a 6-digit numeric OTP using a cryptographically secure source."""
    import secrets
    return str(secrets.randbelow(900000) + 100000)

def send_email_otp(user, otp_code: str):
    """Send OTP to user's email."""
    from django.core.mail import send_mail
    from django.conf import settings
    from django.template.loader import render_to_string
    from django.utils.html import strip_tags

    subject = "Your AI Resume Analyzer Verification Code"
    context = {
        'username': user.username,
        'otp_code': otp_code,
        'valid_minutes': 10,
    }
    
    html_message = render_to_string('analyzer/emails/otp_email.html', context)
    plain_message = strip_tags(html_message)
    
    # Use EMAIL_HOST_USER if configured, else default from address
    from_email = getattr(settings, 'EMAIL_HOST_USER', 'noreply@resumeanalyzer.com')
    if not from_email:
        from_email = 'noreply@resumeanalyzer.com'
    
    send_mail(
        subject,
        plain_message,
        from_email,
        [user.email],
        html_message=html_message,
        fail_silently=False
    )

def send_sms_otp(user, otp_code: str, phone_number: str):
    """Send OTP via SMS using Twilio."""
    from django.conf import settings
    from twilio.rest import Client
    import logging

    logger = logging.getLogger(__name__)

    account_sid = getattr(settings, 'TWILIO_ACCOUNT_SID', None)
    auth_token = getattr(settings, 'TWILIO_AUTH_TOKEN', None)
    twilio_number = getattr(settings, 'TWILIO_PHONE_NUMBER', None)

    if not account_sid or not auth_token or not twilio_number:
        logger.warning(f"Twilio credentials missing. SMS OTP for {user.username} could not be sent.")
        return

    try:
        client = Client(account_sid, auth_token)
        message = client.messages.create(
            body=f"Your Resume Analyzer verification code is: {otp_code}. It expires in 10 minutes.",
            from_=twilio_number,
            to=phone_number
        )
        logger.info(f"Sent SMS to {phone_number}, SID: {message.sid}")
    except Exception as e:
        logger.error(f"Failed to send SMS to {phone_number}: {e}")
        # In a real app, you might want to raise the exception or handle it
        pass

def suggest_bullet_rewrites(bullet_point: str, job_desc: str) -> list:
    """Gets 3 improved options for a resume bullet point from AI."""
    from django.conf import settings
    api_key = getattr(settings, "GROQ_API_KEY", None) or os.environ.get("GROQ_API_KEY", "")
    if not api_key:
        return [bullet_point] * 3
        
    try:
        client = Groq(api_key=api_key)
        prompt = build_bullet_rewrite_prompt(bullet_point, job_desc)
        model = os.environ.get("FAST_MODEL", "llama3-8b-8192")
        
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5,
            max_tokens=500
        )
        
        content = response.choices[0].message.content.strip()
        # Clean any markdown code fences if outputted
        if content.startswith("```json"):
            content = content[7:]
        if content.startswith("```"):
            content = content[3:]
        if content.endswith("```"):
            content = content[:-3]
        content = content.strip()
        
        return json.loads(content)
    except Exception as exc:
        print(f"[utils.suggest_bullet_rewrites] Error: {exc}")
        return [
            f"{bullet_point} (Enhanced version 1 - Quantified outcome)",
            f"{bullet_point} (Enhanced version 2 - Skill integration)",
            f"{bullet_point} (Enhanced version 3 - Strong action verb)"
        ]

def generate_next_interview_question(resume_text: str, job_desc: str, chat_history: list) -> str:
    """Generates the next question in the mock interview chat session."""
    from django.conf import settings
    api_key = getattr(settings, "GROQ_API_KEY", None) or os.environ.get("GROQ_API_KEY", "")
    if not api_key:
        return "Can you tell me about your experience working with technical systems?"
        
    try:
        client = Groq(api_key=api_key)
        prompt = build_interview_question_prompt(resume_text, job_desc, chat_history)
        model = os.environ.get("FAST_MODEL", "llama3-8b-8192")
        
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=300
        )
        return response.choices[0].message.content.strip()
    except Exception as exc:
        print(f"[utils.generate_next_interview_question] Error: {exc}")
        return "Tell me about your background and how it matches this role."

def evaluate_interview_answer(question: str, answer: str, job_desc: str) -> dict:
    """Evaluates candidate response and provides feedback and score."""
    from django.conf import settings
    api_key = getattr(settings, "GROQ_API_KEY", None) or os.environ.get("GROQ_API_KEY", "")
    if not api_key:
        return {"score": 75, "feedback": "Good response. Try to add more metrics."}
        
    try:
        client = Groq(api_key=api_key)
        prompt = build_interview_feedback_prompt(question, answer, job_desc)
        model = os.environ.get("FAST_MODEL", "llama3-8b-8192")
        
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=500
        )
        
        content = response.choices[0].message.content.strip()
        if content.startswith("```json"):
            content = content[7:]
        if content.startswith("```"):
            content = content[3:]
        if content.endswith("```"):
            content = content[:-3]
        content = content.strip()
        
        return json.loads(content)
    except Exception as exc:
        print(f"[utils.evaluate_interview_answer] Error: {exc}")
        return {"score": 70, "feedback": f"Your response is noted. Focus on detailing matching skills. Error: {exc}"}

def parse_resume_to_json(resume_text: str) -> dict:
    """Converts plain text resume to structured JSON format."""
    from django.conf import settings
    api_key = getattr(settings, "GROQ_API_KEY", None) or os.environ.get("GROQ_API_KEY", "")
    if not api_key:
        return {"name": "Candidate", "experience": [], "skills": {}}
        
    try:
        client = Groq(api_key=api_key)
        prompt = build_resume_parser_prompt(resume_text)
        model = os.environ.get("QUALITY_MODEL", "llama-3.3-70b-versatile")
        
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            max_tokens=2000
        )
        
        content = response.choices[0].message.content.strip()
        if content.startswith("```json"):
            content = content[7:]
        if content.startswith("```"):
            content = content[3:]
        if content.endswith("```"):
            content = content[:-3]
        content = content.strip()
        
        return json.loads(content)
    except Exception as exc:
        print(f"[utils.parse_resume_to_json] Error: {exc}")
        # Build basic fallback object by splitting text
        lines = [l.strip() for l in resume_text.split("\n") if l.strip()]
        name = lines[0] if lines else "Candidate"
        return {
            "name": name,
            "contact": {"email": "", "phone": ""},
            "summary": "AI Parsing fell back. Raw text available for editing.",
            "experience": [{"role": "Professional", "company": "Experience", "duration": "", "bullets": [resume_text[:200]]}],
            "education": [],
            "skills": {"other": ["AI parsing failed"]}
        }

def get_ai_summary_suggestions(job_title: str, industry: str) -> list:
    """Invokes AI to get 3 professional summary suggestions."""
    from django.conf import settings
    api_key = getattr(settings, "GROQ_API_KEY", None) or os.environ.get("GROQ_API_KEY", "")
    if not api_key:
        return [
            f"Experienced {job_title} professional in {industry} industry.",
            f"Result-driven professional seeking a {job_title} role to deliver value.",
            f"Passionate specialist with technical competence matching {job_title} requirements."
        ]
    try:
        client = Groq(api_key=api_key)
        prompt = build_summary_suggestion_prompt(job_title, industry)
        model = os.environ.get("FAST_MODEL", "llama3-8b-8192")
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=600
        )
        content = response.choices[0].message.content.strip()
        if content.startswith("```json"):
            content = content[7:]
        if content.startswith("```"):
            content = content[3:]
        if content.endswith("```"):
            content = content[:-3]
        content = content.strip()
        return json.loads(content)
    except Exception as exc:
        print(f"[utils.get_ai_summary_suggestions] Error: {exc}")
        return [
            f"Experienced {job_title} professional in {industry} industry.",
            f"Result-driven professional seeking a {job_title} role to deliver value.",
            f"Passionate specialist with technical competence matching {job_title} requirements."
        ]

def get_ai_experience_bullets(job_title: str, company_type: str) -> list:
    """Invokes AI to get 5 high-impact bullet points."""
    from django.conf import settings
    api_key = getattr(settings, "GROQ_API_KEY", None) or os.environ.get("GROQ_API_KEY", "")
    if not api_key:
        return [f"Delivered key projects for {job_title} at the company."] * 5
    try:
        client = Groq(api_key=api_key)
        prompt = build_experience_bullets_prompt(job_title, company_type)
        model = os.environ.get("FAST_MODEL", "llama3-8b-8192")
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=800
        )
        content = response.choices[0].message.content.strip()
        if content.startswith("```json"):
            content = content[7:]
        if content.startswith("```"):
            content = content[3:]
        if content.endswith("```"):
            content = content[:-3]
        content = content.strip()
        return json.loads(content)
    except Exception as exc:
        print(f"[utils.get_ai_experience_bullets] Error: {exc}")
        return [
            f"Spearheaded key initiatives as {job_title} contributing to team success.",
            f"Optimized system performance and architecture leading to tangible improvements.",
            f"Collaborated with cross-functional partners to execute on product requirements.",
            f"Successfully designed and implemented services using industry best practices.",
            f"Resolved critical issues under tight deadlines ensuring high availability."
        ]

