import os
from docx import Document

fixtures_dir = os.path.join(os.path.dirname(__file__), "fixtures", "ats_hostile")
os.makedirs(fixtures_dir, exist_ok=True)

# 1. Clean control
doc = Document()
doc.add_paragraph("Ayan Bhatia")
doc.add_paragraph("ayan@example.com | 555-123-4567")
doc.add_paragraph("EXPERIENCE")
doc.add_paragraph("Software Engineer")
doc.add_paragraph("2020 - Present")
doc.save(os.path.join(fixtures_dir, "clean_control.docx"))

# 2. Multi-column table
doc = Document()
doc.add_paragraph("Ayan Bhatia")
doc.add_paragraph("ayan@example.com | 555-123-4567")
doc.add_paragraph("EXPERIENCE")
table = doc.add_table(rows=1, cols=2)
row = table.rows[0]
row.cells[0].text = "2020 - Present"
row.cells[1].text = "Software Engineer"
doc.save(os.path.join(fixtures_dir, "multi_column_table.docx"))

# Note: We can't easily generate all hostile features (like header/footer, hidden text, graphics) 
# reliably using python-docx without complex manipulation, but we can test the ones we detect natively.
# Currently, `extract_text` in utils.py only checks `len(doc.tables) > 0` for DOCX.
# For PDFs it checks PyMuPDF `find_tables` and `get_images`.

print("Fixtures generated successfully.")
